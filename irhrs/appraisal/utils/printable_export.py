from docx import Document
from html2text import html2text

from irhrs.appraisal.models.appraiser_setting import Appraisal
from irhrs.appraisal.models.performance_appraisal_setting import ScoreAndScalingSetting
from irhrs.questionnaire.models.helpers import (
    LINEAR_SCALE, RATING_SCALE, LONG, SHORT, DATE, TIME, DURATION, DATE_TIME,
    DATE_WITHOUT_YEAR, DATE_TIME_WITHOUT_YEAR, FILE_UPLOAD,
    RADIO, CHECKBOX
)
from irhrs.core.utils import nested_get

def generate_printable_document(appraisal: Appraisal) -> Document:
    # Initialize Document for Export
    document = Document()
    # Page Title = question_set.title & description
    document.add_heading(appraisal.question_set['title'], 0)
    document.add_paragraph(html2text(appraisal.question_set['description']))

    question_set = appraisal.question_set  # JSONField()
    for section in question_set['sections']:
        # for each section, add heading and subheading
        document.add_heading(section['title'], 1)
        document.add_heading(section['description'], 2)
        is_feedback_section = section['title'] == 'Feedback'

        for question_number, question in enumerate(section['questions'], start=1):

            title = question['question']['title']
            question_description = question['question']['description']
            remarks = question['question']['remarks']
            if is_feedback_section:
                stringify_question_number = ''
            else:
                stringify_question_number = '%s. ' % str(question_number)
            clean_title = stringify_question_number + html2text(title)

            is_open_ended = question['question']['is_open_ended']
            answer_choices = question['question']['answer_choices']
            rating_scale = question['question']['rating_scale']
            is_mandatory = question.get('is_mandatory')

            paragraph = document.add_paragraph('\n')
            if is_mandatory:
                paragraph.add_run('*').bold = True
            paragraph.add_run(clean_title)
            if question_description:
                paragraph.add_run(question_description)

            # initial guess for letter size (todo @rawV: verify in ms-docx)
            page_width = 78  # 78 + '\n' = 79 in total

            choices = nested_get(question, 'question.answers')
            if answer_choices in [LINEAR_SCALE, RATING_SCALE]:
                para = document.add_paragraph()
                table = document.add_table(rows=len(choices) + 1, cols=3)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'SN'
                table.cell(0, 1).text = 'Name'
                table.cell(0, 2).text = 'Remarks'

                for index, choice in enumerate(choices, start=1):
                    if not choice:
                        continue
                    remarks_cell = table.cell(index, 2)
                    name_cell = table.cell(index, 1)
                    sn_cell = table.cell(index, 0)

                    sn_cell.text = str(index)
                    name_cell.text = choice['name']
                    remarks_cell.text = remarks

            elif answer_choices == LONG:
                document.add_paragraph(choices)

            elif answer_choices == SHORT:
                document.add_paragraph(choices)

            elif answer_choices in [
                DATE,
                TIME,
                DURATION,
                DATE_TIME,
                DATE_WITHOUT_YEAR,
                DATE_TIME_WITHOUT_YEAR,
                FILE_UPLOAD,
            ]:
                document.add_paragraph('_' * int(page_width / 4))
            elif answer_choices in (CHECKBOX, RADIO):
                # generic message before a MCQ table
                para = document.add_paragraph()
                choices = choices
                # choices = {'id': 313, 'order': 0, 'title': 'Option C', 'remarks': '', 'is_correct': False}

                table = document.add_table(rows=len(choices) + 1, cols=3)
                table.style = 'Table Grid'
                table.cell(0, 0).text = 'SN'
                table.cell(0, 1).text = 'Answer'
                table.cell(0, 2).text = 'Remarks'

                for index, choice in enumerate(choices, start=1):
                    sn_cell = table.cell(index, 0)
                    answer_cell = table.cell(index, 1)
                    remarks_cell = table.cell(index, 2)

                    sn_cell.text = str(index)
                    answer_cell.text = choice['title']
                    remarks_cell.text = choice['remarks']

            else:
                # MULTIPLE_CHOICE_GRID and CHECKBOX_GRID is unsupported for now
                document.add_paragraph('_' * page_width)
                document.add_paragraph('_' * page_width)
                document.add_paragraph('_' * page_width)

            """
                ----------------> (table)
                [ ] MULTIPLE_CHOICE_GRID = 'multiple-choice-grid'
                [ ] CHECKBOX_GRID = 'checkbox-grid'
                [X] CHECKBOX = 'multiple-mcq'
                [X] RADIO = 'single-mcq'

                --------------->(table)
                [X] LINEAR_SCALE = 'linear-scale'
                [X] RATING_SCALE = 'rating-scale'

                ---------------->(page-width * '_' )
                [X] SHORT = 'short-text'
                [X] LONG = 'long-text'

                ---------------->(page-width / 4 * '_' )

                [X] DATE = 'date'
                [X] TIME = 'time'
                [X] DURATION = 'duration'
                [X] DATE_TIME = 'date-time'
                [X] DATE_WITHOUT_YEAR = 'date-without-year'
                [X] DATE_TIME_WITHOUT_YEAR = 'date-time-without-year'
                [X] FILE_UPLOAD = 'file-upload'
            """
    return document
